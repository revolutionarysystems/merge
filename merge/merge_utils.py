import os
import io
import re
import string
import zipfile
import tempfile
import shutil
from docx import Document
#  from docx.text.paragraph import Paragraph
from django.template import Template, Context
from markdown import markdown
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import difflib
from subprocess import check_output
import lxml.etree as etree

from .resource_utils import (
    get_working_dir, strip_xml_dec, get_xml_dec, get_output_dir, get_local_dir, get_local_txt_content)
from .config import install_name

def replaceParams(txt, subs):
    for key in subs.keys():
        old = "${"+key+"}"
        if (txt.find(old)>=0):
            txt = txt.replace(old, subs[key])
    return txt

def removePara(para):
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None

def substituteVariablesPlain(fileNameIn, fileNameOut, subs):
    c = Context(subs)
    fileIn = open(fileNameIn, "r")
    fullText = fileIn.read()
    t = Template(fullText)
    xtxt = t.render(c)
    fileOut = open(fileNameOut, "w")
    fileOut.write(xtxt)
    return {"file":fileNameOut}
    
def substituteVariablesPlainString(stringIn, subs):
    c = Context(subs)
    fullText = stringIn
    t = Template(fullText)
    xtxt = t.render(c)
    return xtxt
    
def preprocess(text):
    text = text.replace("{% #A", "{% cycle 'A' 'B' 'C' 'D' 'E' 'F' 'G' 'H' 'I' 'J' 'K' 'L' 'M' 'N' 'O' 'P' 'Q' 'R' 'S' 'T' 'U' 'V' 'W' 'X' 'Y' 'Z'")    
    text = text.replace("{% #a", "{% cycle 'a' 'b' 'c' 'd' 'e' 'f' 'g' 'h' 'i' 'j' 'k' 'l' 'm' 'n' 'o' 'p' 'q' 'r' 's' 't' 'u' 'v' 'w' 'x' 'y' 'z'")    
    text = text.replace("{% #9", "{% cycle '1' '2' '3' '4' '5' '6' '7' '8' '9' '10' '11' '12' '13' '14' '15' '16' '17' '18' '19' '20' '21' '22' '23' '24' '25' '26' '27' '28' '29' '30' '31' '32' '33' '34' '35' '36' '37' '38' '39' '40' '41' '42' '43' '44' '45' '46' '47' '48' '49'")    
    text = text.replace("{% #I", "{% cycle 'I' 'II' 'III' 'IV' 'V' 'VI' 'VII' 'VIII' 'IX' 'X' 'XI' 'XII' 'XIII' 'XIV' 'XV' 'XVI' 'XVII' 'XVIII' 'XIX' 'XX'")    
    text = text.replace("{% #i", "{% cycle 'i' 'ii' 'iii' 'iv' 'v' 'vi' 'vii' 'viii' 'ix' 'x' 'xi' 'xii' 'xiii' 'xiv' 'xv' 'xvi' 'xvii' 'xviii' 'xix' 'xx'")    
    return text

def apply_sequence(text):
    alf = string.ascii_uppercase
    target = '[% #A %]'
    sub = text.find(target)
    n= 0
    while sub>=1:
        text = text[:sub]+alf[n]+text[sub+len(target):]
        n+=1
        sub = text.find(target)
    return text

def docx_copy_run_style_from(run1, run2):
    run1.font.color.rgb = run2.font.color.rgb
    run1.font.all_caps = run2.font.all_caps
    run1.font.bold = run2.font.bold
    run1.font.italic = run2.font.italic
    run1.font.size = run2.font.size
    run1.font.underline = run2.font.underline
    #complex_script, cs_bold, cs_italic, double_strike, emboss, hidden, highlight_color,
    #imprint, math, name, no_proof, outline, rtl, shadow, small_caps, snap_to_grid, spec_vanish, 
    #strike, superscript, underline, web_hidden

def docx_copy_para_format_from(para1, para2):
    para1.paragraph_format.alignment = para2.paragraph_format.alignment
    para1.paragraph_format.first_line_indent = para2.paragraph_format.first_line_indent
    para1.paragraph_format.keep_together = para2.paragraph_format.keep_together
    para1.paragraph_format.keep_with_next = para2.paragraph_format.keep_with_next
    para1.paragraph_format.left_indent = para2.paragraph_format.left_indent
    try:
        para1.paragraph_format.line_spacing = para2.paragraph_format.line_spacing
        para1.paragraph_format.line_spacing_rule = para2.paragraph_format.line_spacing_rule
    except ValueError:
        pass
    para1.paragraph_format.page_break_before = para2.paragraph_format.page_break_before
    para1.paragraph_format.right_indent = para2.paragraph_format.right_indent
    para1.paragraph_format.space_after = para2.paragraph_format.space_after
    para1.paragraph_format.space_before = para2.paragraph_format.space_before
    para1.paragraph_format.widow_control = para2.paragraph_format.widow_control

def isControlText(s):
    s = s.strip()
    if s[:2]=="{%" and s[-2:]=="%}" and s.find("%}")== s.rfind("%}"):
        if s.find("include")>=0:
            return False
        elif s.find("now")>=0:
            return False
        else:
            return True
    else:
        return False

def isControlLine(s):
    s = s.split("+")[0]
    return isControlText(s)
#    s = s.strip()
#    if s[:2]=="{%" and s[-2:]=="%}" and s.find("%}")== s.rfind("%}"):
#        if s.find("include")>=0:
#            return False
#        else:
#            return True
#    else:
#        return False

def wrap_list_xml(childlist, root_tag, child_tag):
    out = "<"+root_tag+">\n"
    for child in childlist:
        out+= "\t<"+child_tag+">"+child+"</"+child_tag+">\n"
    out+= "</"+root_tag+">"
    return out

def docx_text(file_name_in):
    doc_in = Document(docx=file_name_in)
    paras=doc_in.paragraphs
    fullText="" 
    for para in paras:
        paraText=para.text
        fullText+=paraText
    return fullText

def extract_regex_matches_docx(file_name_in, regex, wrap=None, root_tag="list", child_tag="item"):
    text = docx_text(file_name_in)
    p = re.compile(regex)
    m=p.findall(text)
    if wrap:
        return wrap_list_xml(m, root_tag, child_tag)
    else:
        return m

    #            literal = m.group('literal')


def preprocess_docx_template(file_name_in, file_name_out):
    timestamp_template = os.stat(file_name_in).st_mtime
    prep_file_exists = False
    if os.path.isfile(file_name_out):
        timestamp_template_prep = os.stat(file_name_out).st_mtime
        prep_file_exists = True
    if (not prep_file_exists) or timestamp_template_prep < timestamp_template:
        doc_in = Document(docx=file_name_in)
        paras=doc_in.paragraphs
        i = 0
        for para in paras:
            paraText=""
            runs = para.runs
            for run in runs:
                txt = run.text
                if isControlText(txt):
                    run.text = "[##]"+txt
        doc_in.save(file_name_out)

def postprocess_docx(file_name_in):
    doc_in = Document(docx=file_name_in)
    for p in doc_in.paragraphs:
        if p.text.strip()=="[##]":
            removePara(p)
    doc_in.save(file_name_in)


def substituteVariablesDocx(file_name_in, fileNameOut, subs):
    c = Context(subs)
    doc_in = Document(docx=file_name_in)
    doc_temp = Document()
    paras=doc_in.paragraphs
    fullText="" 
    i = 0
    styles = {}
    for para in paras:
        paraText=""
        p = doc_temp.add_paragraph(style = para.style)
        docx_copy_para_format_from(p, para)
        j = 0
        runs = para.runs
        for run in runs:
            txt = run.text
            paraText+= txt+"+"+str(j)+"+run+"
            r = p.add_run(text = txt, style=run.style)
            docx_copy_run_style_from(r, run)
            j+=1
        fullText+= paraText+str(i)+"+para+"
        i+=1
    fullText = preprocess(fullText)
    t = Template(fullText)
    xtxt = t.render(c)
    xtxt = apply_sequence(xtxt)
    xParaTxts = xtxt.split("+para+")
    for p in paras:
        removePara(p)

    doc_in.paragraphs.clear()
    paras=doc_temp.paragraphs
    for xParaTxt in xParaTxts:
        runTxts = xParaTxt.split("+run+")
        if runTxts[-1]!='':
            para_n = int(runTxts[-1])
            p = doc_in.add_paragraph(style=paras[para_n].style)
            docx_copy_para_format_from(p, paras[para_n])
            for runTxt in runTxts[:-1]:
                try:
                    txt = runTxt.split("+")[-2]
                except:
                    txt=""
                run_n = int(runTxt.split("+")[-1])
                r = p.add_run(text=txt, style=paras[para_n].runs[run_n].style)
                docx_copy_run_style_from(r, paras[para_n].runs[run_n])
            if isControlLine(paras[para_n].text):
                p.text="{}"

    for p in doc_in.paragraphs:
        if p.text=="{}":
            removePara(p)
    doc_in.save(fileNameOut)
    return {"file":fileNameOut}


def print_doc(doc):
    paras=doc.paragraphs
    print("...")
    for para in paras[14:20]:
        print(para.text)
    print("...")

def combine_docx(file_names, file_name_out):
    combined_document = Document(file_names[0])
    count, number_of_files = 0, len(file_names)
    for file in file_names[1:]:
        if file == "pagebreak":
            combined_document.add_page_break()
        else:    
            sub_doc = Document(file)
            for para in sub_doc.paragraphs:
                pnew = combined_document.add_paragraph(style=para.style)
                docx_copy_para_format_from(pnew, para)
                runs = para.runs
                for run in runs:
                    rnew = pnew.add_run(text=run.text, style=run.style)
                    docx_copy_run_style_from(rnew, run)

    combined_document.save(file_name_out)
    return {"file":file_name_out}

def convert_markdown(fileNameIn, fileNameOut):
    fileIn = open(fileNameIn, "r")
    fileOut = open(fileNameOut, "w")
    fileOut.write(markdown(fileIn.read()))
    fileOut.close()
    return {"file":fileNameOut}

def email_file(baseFileName, me, you, subject, credentials):

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = me
    msg['To'] = you.replace(" ","+")

    mimetypes = ["text/plain", "text/html"]
    file_exts = [".md", ".html"]

    for mime in zip(mimetypes, file_exts):
        file=baseFileName+mime[1]
        fp = open(file, 'r')
        content = fp.read()
        fp.close()
#        print(mime[0].split("/"))
        msg.attach(MIMEText(content, mime[0].split("/")[1]))    
    
    username = credentials["username"]
    password = credentials["password"]
    server = smtplib.SMTP(credentials["server"])

    server.ehlo()
    server.starttls()
    server.ehlo()
    server.login(username,password)
#    print(me)
#    print(you.replace(" ","+"))
#    print(msg.as_string())
    response = server.sendmail(me, [you.replace(" ","+")], msg.as_string())
    server.quit()
    return {"email":you.replace(" ","+")}

def merge_docx_footer(full_local_filename, subs):
    merge_docx_header_footer(full_local_filename, subs, "footer1")
    return merge_docx_header_footer(full_local_filename, subs, "footer2")

def merge_docx_header(full_local_filename, subs):
    merge_docx_header_footer(full_local_filename, subs, "header1")
    return merge_docx_header_footer(full_local_filename, subs, "header2")

def merge_docx_header_footer(full_local_filename, subs, xmlname):
    docx_filename = full_local_filename
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    xml_content = zip.read('word/'+xmlname+'.xml')
 #   xml_content = xml_content.decode("ISO-8859-1")
###
    xml_content = xml_content.decode("UTF-8")
    try:
        xml_content = substituteVariablesPlainString(xml_content, subs)
    except:
        pass
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)

    with open(os.path.join(tmp_dir,'word/'+xmlname+'.xml'), 'w') as f:
        f.write(xml_content)
    filenames = zip.namelist()
    zip_copy_filename = docx_filename
    with zipfile.ZipFile(zip_copy_filename, "w") as docx:
        for filename in filenames:
            docx.write(os.path.join(tmp_dir,filename), filename)
    shutil.rmtree(tmp_dir)
    return({"file":docx_filename})

def docx_subfile(zip, tmp_dir, subs, filename):
    try:
        xml_content = zip.read(filename)
        xml_content = xml_content.decode("UTF-8")
        #xml_content = xml_content.decode("ISO-8859-1")
        xml_content = preprocess(xml_content)
        xml_content = xml_content.replace("&quot;", '"')
        xml_content_subs = str(substituteVariablesPlainString(xml_content, subs))
        with io.open(os.path.join(tmp_dir,filename), 'w', encoding="UTF-8") as f:
#        with io.open(os.path.join(tmp_dir,filename), 'w', encoding="UTF-8") as f:
            f.write(xml_content_subs)
    except KeyError:
        pass

def substituteVariablesDocx_direct(file_name_in, file_name_out, subs):


    docx_filename = file_name_in
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    filenames = zip.namelist()
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)

    candidates = ["word/document.xml","word/header1.xml","word/header2.xml","word/footer1.xml","word/footer2.xml"]
    for filename in filenames:
        if filename in candidates:
            docx_subfile(zip, tmp_dir, subs, filename)

    with zipfile.ZipFile(file_name_out, "w") as docx:
        for filename in filenames:
            docx.write(os.path.join(tmp_dir,filename), filename)
    shutil.rmtree(tmp_dir)

    return({"file":file_name_out})

def get_docx_paras(zip):
    xml_content = zip.read("word/document.xml").decode("utf8")
    paras_start = xml_content.find("<w:p")
    paras_end = xml_content.rfind("</w:p>")+6
    return xml_content[paras_start:paras_end]

def get_docx_numbering(zip):
    xml_content = zip.read("word/numbering.xml").decode("utf8")
    remain = xml_content
    abstract_num_schemes = {}
    num_schemes = {}
    while remain.find("<w:abstractNum ")>=0:
        scheme_start = remain.find("<w:abstractNum ")
        scheme_end = remain.find("</w:abstractNum>")+16
        scheme_str = remain[scheme_start:scheme_end]
        scheme_index_loc_s = scheme_str.find('="')
        scheme_index_loc_e = scheme_str.find('">')
        scheme_index = scheme_str[scheme_index_loc_s+2:scheme_index_loc_e]
        abstract_num_schemes[scheme_index] = scheme_str
        remain = remain[scheme_end:]
    while remain.find("<w:num ")>=0:
        scheme_start = remain.find("<w:num ")
        scheme_end = remain.find("</w:num>")+8
        scheme_str = remain[scheme_start:scheme_end]
        scheme_index_loc_s = scheme_str.find('="')
        scheme_index_loc_e = scheme_str.find('">')
        scheme_index = scheme_str[scheme_index_loc_s+2:scheme_index_loc_e]
        num_schemes[scheme_index] = scheme_str
        remain = remain[scheme_end:]
    return abstract_num_schemes, num_schemes, xml_content


def get_docx_content(filename):
    docx_filename = filename
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    filenames = zip.namelist()
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)
    xml_str = zip.read("word/_rels/document.xml.rels").decode("utf8")
    rel_xml_content = strip_xml_dec(xml_str)
    rel_xml_dec = get_xml_dec(xml_str)
    rel_dom = etree.fromstring(rel_xml_content)
    response = {}
    response["paras"] = get_docx_paras(zip)
    response["numbering"] = get_docx_numbering(zip)
    response["tmp_dir"] = tmp_dir
    response["filenames"] = filenames
    response["rel_dom"] = rel_dom
    # get images too
    # return tmp_dir too
    return response

def copy_docx_media(tmp_dir_from, tmp_dir_to, filenames_from, filenames_to, rel_0, max_Rel_id):
    rel_elements = []
    renames = []
    for filename in filenames_from:
        if filename.find("word/media/")==0:
#            filename=filename.replace("/",os.sep)
            print("media file:",tmp_dir_from, tmp_dir_to,  filename)
            source = os.path.join(tmp_dir_from,filename)
            dest = os.path.join(tmp_dir_to,filename)
            if not os.path.exists(os.path.split(dest)[0]):
                os.makedirs(os.path.split(dest)[0])
            print(filenames_to)
            if filename in filenames_to:
                oldfilename = filename
                i = 1
                while filename in filenames_to:
                    filename = "word/media/image"+"{0:0>2}".format(i)+".png"
                    i+=1
                print(oldfilename,">",filename)
                renames.append((os.path.split(oldfilename)[1],os.path.split(filename)[1]))
                dest = os.path.join(tmp_dir_to,filename)
            else:
                renames.append((os.path.split(filename)[1],os.path.split(filename)[1]))
            shutil.copyfile(source.replace("/",os.sep),dest.replace("/",os.sep))
            filenames_to.append(filename)
            new_rel_el = etree.Element("Relationship", attrib={"Id":"".join(["rId", str(max_Rel_id+1)]), 
                "Type":rel_0.attrib["Type"].replace("settings","image"), 
                "Target":"/".join(["media",os.path.split(filename)[1]])})
            rel_elements.append(new_rel_el)
            max_Rel_id+=1
    return filenames_to, rel_elements, renames

def combine_docx_direct(file_names_to_combine, file_name_out):
    docx_filename = file_names_to_combine[0]
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    filenames = zip.namelist()
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)
    main_xml_content = zip.read("word/document.xml").decode("utf8")
    xml_str = zip.read("word/_rels/document.xml.rels").decode("utf8")
    rel_xml_content = strip_xml_dec(xml_str)
    rel_xml_dec = get_xml_dec(xml_str)
    rel_dom = etree.fromstring(rel_xml_content)
    max_Rel_id = len(rel_dom)
    rel_0  = rel_dom[0]
    main_abs_num, main_num, main_number_xml = get_docx_numbering(zip)
    insertion_point = main_xml_content.find("<w:sectPr>")
    count, number_of_files = 0, len(file_names_to_combine)
    for file in file_names_to_combine[1:]:
        if file == "pagebreak":
            break_xml = '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'
            main_xml_content = main_xml_content[:insertion_point]+break_xml+main_xml_content[insertion_point:]
            insertion_point+=len(break_xml)
        else:    
            print()
            print("next file:", file)
            sub_xml_content = get_docx_content(file)
            sub_xml_paras = sub_xml_content["paras"]
            abs_num, num, sub_num_xml = sub_xml_content["numbering"]
            if not main_abs_num: # No main scheme
                if abs_num: # But there is sub-document numbering
                    main_number_xml = sub_num_xml
                    main_abs_num = abs_num
                    main_num = num
            else: # there is a main scheme
                if abs_num: # And there is sub-document numbering
                    for key in abs_num.keys():
                        n_main_schemes = len(main_abs_num)
                        new_key = str(n_main_schemes+1)
                        main_abs_num[new_key]=abs_num[key].replace('w:abstractNumId="'+key+'"', 'w:abstractNumId="'+new_key+'"')
                        main_num[new_key]=num[key].replace('w:numId="'+key+'"', 'w:numId="'+new_key+'"').replace('<w:abstractNumId w:val="'+key+'"/>','<w:abstractNumId w:val="'+new_key+'"/>')
                        sub_xml_paras = sub_xml_paras.replace('<w:numId w:val="'+key+'"/>', '<w:numId w:val="'+new_key+'"/>')


            #Now move media files:  # incomplete - also need to correct relationships
            filenames, rel_elements, renames= copy_docx_media(sub_xml_content["tmp_dir"], tmp_dir, sub_xml_content["filenames"], filenames, rel_0, max_Rel_id)

#            rel_renames = []
            for rename in renames:
                print(rename)
                sub_xml_paras = sub_xml_paras.replace(rename[0],rename[1])
                for element in rel_elements:
                    if element.attrib["Target"]=="media/"+rename[1]:
                        new_rel = element.attrib["Id"]

                        # make sure the replacement is right

                for element in sub_xml_content["rel_dom"]:
                    if element.attrib["Target"]=="media/"+rename[0]:
                        old_rel = element.attrib["Id"]
                sub_xml_paras = sub_xml_paras.replace(old_rel,new_rel)
#                rel_renames.append((old_rel, new_rel))

 
            for element in rel_elements:
                rel_dom.append(element)

            max_Rel_id += len(rel_elements)

            main_xml_content = main_xml_content[:insertion_point]+sub_xml_paras+main_xml_content[insertion_point:]
            insertion_point+=len(sub_xml_paras)
            if main_number_xml.find("<w:abstractNum ") >=0:
                main_number_xml_pre = main_number_xml[:main_number_xml.find("<w:abstractNum ")]
                main_number_xml_post = main_number_xml[main_number_xml.rfind("</w:num>")+8:]
                main_number_xml = main_number_xml_pre
                for scheme in main_abs_num.values():
                    main_number_xml+=scheme
                for scheme in main_num.values():
                    main_number_xml+=scheme
                main_number_xml+=main_number_xml_post
            #print(main_number_xml)
    rels_xml = "\n".join([rel_xml_dec,etree.tostring(rel_dom).decode("utf8")])
    print(rels_xml)



    #All files assimilated
    with io.open(os.path.join(tmp_dir,"word/document.xml"), 'w', encoding="utf8") as f:
        f.write(main_xml_content)
    with io.open(os.path.join(tmp_dir,"word/numbering.xml"), 'w', encoding="utf8") as f:
        f.write(main_number_xml)
    with io.open(os.path.join(tmp_dir,"word/_rels/document.xml.rels"), 'w', encoding="utf8") as f:
        f.write(rels_xml)
    # rewrite the rel file
    with zipfile.ZipFile(file_name_out, "w") as docx:
        for filename in filenames:
            docx.write(os.path.join(tmp_dir,filename), filename)
    shutil.rmtree(tmp_dir)
    return({"file":file_name_out})

def shellCommand(command):
    return check_output(command, shell=True).decode()

def convert_pdf(filename_in, filename_out, outdir = "."):
    command = "soffice --headless --convert-to pdf "+filename_in+" --outdir "+outdir
    response = shellCommand("soffice --headless --convert-to pdf "+filename_in+" --outdir "+outdir)
    return {"file":filename_out, "response":response, "command": command}
